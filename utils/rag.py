import os
import shutil
import fitz  # PyMuPDF
import docx
from typing import List
from pathlib import Path

from langchain_community.vectorstores import FAISS
from langchain.schema import Document
from langchain_ollama import OllamaLLM
from langchain_openai import OpenAIEmbeddings
from langdetect import detect as lang_detect

# EmbeddingType enum simplifié
class EmbeddingType:
    LOCAL = "local"
    OPENAI = "openai"
    HUGGINGFACE = "huggingface"

llm = OllamaLLM(model="llama3.2:1b", base_url="http://localhost:11434")

INDEX_DIR = "faiss_index"

def clean_index(folder: str):
    if os.path.exists(folder):
        shutil.rmtree(folder)
    os.makedirs(folder)

def load_documents(file_paths: List[str]) -> List[Document]:
    """
    Charge et lit plusieurs fichiers (.txt, .pdf, .docx) directement.
    Chaque chemin dans file_paths doit être un fichier, pas un dossier.
    """
    docs = []
    for path in file_paths:
        filename = os.path.basename(path)
        try:
            content = ""
            if filename.endswith(".txt"):
                with open(path, "r", encoding="utf-8") as f:
                    content = f.read()
            elif filename.endswith(".pdf"):
                with fitz.open(path) as doc:
                    for page in doc:
                        content += page.get_text()
            elif filename.endswith(".docx"):
                docx_file = docx.Document(path)
                content = "\n".join([para.text for para in docx_file.paragraphs])
            else:
                print(f"[⚠] Format non supporté : {filename}")
                continue

            if content.strip():
                try:
                    lang = lang_detect(content[:500])
                except:
                    lang = "unknown"

                docs.append(Document(
                    page_content=content,
                    metadata={
                        "source": filename,
                        "lang": lang,
                        "file_path": path
                    }
                ))
        except Exception as e:
            print(f"[❌] Erreur fichier {filename} : {e}")
    return docs


def build_vectorstore_from_files(
    files: List[str],
    index_folder: str,
    embedding_type: EmbeddingType = EmbeddingType.LOCAL
) -> FAISS:
    print("[⚙️] Nettoyage et création de l’index...")
    clean_index(index_folder)

    documents = load_documents(files)  # ⚡ Liste de fichiers directement
    if not documents:
        print(f"[⚠] Aucun document trouvé dans {files}")
        return None

    try:
        if embedding_type == EmbeddingType.LOCAL:
            from langchain_ollama import OllamaEmbeddings
            embedding_model = OllamaEmbeddings(model="llama3.2:1b")
        elif embedding_type == EmbeddingType.OPENAI:
            embedding_model = OpenAIEmbeddings()
        elif embedding_type == EmbeddingType.HUGGINGFACE:
            from langchain_community.embeddings import HuggingFaceEmbeddings
            embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        else:
            raise ValueError(f"Type non supporté : {embedding_type}")

        vectorstore = FAISS.from_documents(documents, embedding_model)
        vectorstore.save_local(index_folder)
        print(f"[✅] Index sauvegardé dans {index_folder}")
        return vectorstore

    except Exception as e:
        print(f"[❌] Échec indexation : {e}")
        return None


def query_rag(question: str, vectorstore: FAISS, user_emotion: str = "neutre", k: int = 5, score_threshold: float = 0.7) -> str:
    """
    user_emotion : émotion détectée (ex. "joyeux", "colère", "stressé", "fatigué", etc.)
    """
    try:
        user_lang = lang_detect(question)
        print(f"[🌐] Langue détectée: {user_lang} | Émotion détectée: {user_emotion}")

        docs_with_scores = vectorstore.similarity_search_with_score(question, k=k)
        
        relevant_docs = []
        for doc, score in docs_with_scores:
            doc_lang = doc.metadata.get("lang", "unknown")
            if score > score_threshold and doc_lang == user_lang:
                relevant_docs.append(doc)
                print(f"[📄] Document sélectionné: {doc.metadata['source']} (score: {score:.2f}, langue: {doc_lang})")
        
        if not relevant_docs:
            return f"Je n'ai pas trouvé d'informations pertinentes pour répondre à votre question. (Émotion détectée : {user_emotion})"
        
        context = "\n\n".join([doc.page_content for doc in relevant_docs])

        # On ajoute l'émotion dans le prompt
        if user_lang.startswith("en"):
            prompt = f"""You are an expert assistant. 
The user is currently feeling: {user_emotion}.
Answer the question considering the user's emotional state, while using ONLY the context below.
If you don't know the answer, say you don't know. Be precise and factual.

Context:
{context}

Question:
{question}

Answer:"""
        elif user_lang.startswith("ar"):
            prompt = f"""أنت مساعد خبير. 
المستخدم حالياً يشعر بـ: {user_emotion}.
أجب على السؤال مع مراعاة الحالة العاطفية للمستخدم، باستخدام السياق أدناه فقط.
إذا كنت لا تعرف الإجابة، فقل أنك لا تعرف. كن دقيقًا وواقعيًا.

السياق:
{context}

السؤال:
{question}

الإجابة:"""
        else:
            prompt = f"""Tu es un assistant expert. 
L'utilisateur se sent actuellement : {user_emotion}.
Réponds à la question en tenant compte de son état émotionnel, en utilisant UNIQUEMENT le contexte ci-dessous.
Si tu ne connais pas la réponse, dis que tu ne sais pas. Sois précis et factuel.

Contexte:
{context}

Question:
{question}

Réponse:"""

        answer = llm.invoke(prompt)
        return answer.strip()

    except Exception as e:
        return f"❌ Erreur RAG : {e}"
